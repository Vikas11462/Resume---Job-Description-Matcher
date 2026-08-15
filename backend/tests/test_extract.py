"""Unit tests for document extraction engine (app.extract)."""

from pathlib import Path
import docx
import pytest
from app.extract import extract_text, extract_text_from_docx, extract_text_from_pdf


def generate_minimal_valid_pdf(output_path: Path, sample_text: str = "Vikas B.Tech CSE Python Developer") -> Path:
    """Generates a minimal, syntactically valid PDF 1.4 binary file for testing."""
    stream_content = f"BT\n/F1 12 Tf\n50 750 Td\n({sample_text}) Tj\nET".encode("latin-1")
    stream_len = len(stream_content)

    obj4 = f"4 0 obj << /Length {stream_len} >> stream\n".encode("latin-1") + stream_content + b"\nendstream\nendobj\n"

    obj1 = b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
    obj2 = b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
    obj3 = b"3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >> endobj\n"
    obj5 = b"5 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"

    header = b"%PDF-1.4\n"
    body = header

    offsets = [0]
    for obj in [obj1, obj2, obj3, obj4, obj5]:
        offsets.append(len(body))
        body += obj

    xref_offset = len(body)
    xref = f"xref\n0 {len(offsets)}\n0000000000 65535 f \n".encode("latin-1")
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode("latin-1")

    trailer = f"trailer << /Size {len(offsets)} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("latin-1")

    full_pdf_bytes = body + xref + trailer
    output_path.write_bytes(full_pdf_bytes)
    return output_path


def generate_sample_docx(output_path: Path) -> Path:
    """Generates a sample DOCX resume document for testing."""
    doc = docx.Document()
    doc.add_heading("Vikas - Software Engineer", level=1)
    doc.add_paragraph("Core competencies: Python, FastAPI, React, Next.js, and Machine Learning.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Education"
    table.cell(0, 1).text = "B.Tech CSE, Jaypee University"
    table.cell(1, 0).text = "Experience"
    table.cell(1, 1).text = "Full Stack NLP Development"
    doc.save(str(output_path))
    return output_path


def test_extract_valid_docx(tmp_path: Path):
    """Verifies text extraction from a valid Microsoft Word document (.docx)."""
    docx_file = tmp_path / "sample_resume.docx"
    generate_sample_docx(docx_file)

    extracted_text = extract_text(docx_file)

    assert "Vikas - Software Engineer" in extracted_text
    assert "FastAPI" in extracted_text
    assert "Jaypee University" in extracted_text


def test_extract_valid_pdf(tmp_path: Path):
    """Verifies text extraction from a valid PDF document (.pdf)."""
    pdf_file = tmp_path / "sample_resume.pdf"
    generate_minimal_valid_pdf(pdf_file, "Vikas B.Tech CSE Python Developer")

    extracted_text = extract_text(pdf_file)

    assert "Vikas" in extracted_text
    assert "Python Developer" in extracted_text


def test_extract_unsupported_extension(tmp_path: Path):
    """Verifies that attempting to parse unsupported formats (e.g. .txt, .png) raises ValueError."""
    txt_file = tmp_path / "notes.txt"
    txt_file.write_text("Plain text content")

    with pytest.raises(ValueError, match="Unsupported file format"):
        extract_text(txt_file)


def test_extract_nonexistent_file():
    """Verifies that non-existent paths raise FileNotFoundError."""
    missing_file = Path("non_existent_resume_file.pdf")

    with pytest.raises(FileNotFoundError, match="File not found"):
        extract_text(missing_file)


def test_extract_corrupted_pdf(tmp_path: Path):
    """Verifies that corrupted PDF documents raise a descriptive ValueError."""
    corrupt_pdf = tmp_path / "corrupted.pdf"
    corrupt_pdf.write_bytes(b"NOT A REAL PDF FILE CONTENT AT ALL")

    with pytest.raises(ValueError, match="Failed to parse PDF"):
        extract_text(corrupt_pdf)


def test_extract_corrupted_docx(tmp_path: Path):
    """Verifies that corrupted DOCX documents raise a descriptive ValueError."""
    corrupt_docx = tmp_path / "corrupted.docx"
    corrupt_docx.write_bytes(b"NOT A REAL DOCX ZIP ARCHIVE")

    with pytest.raises(ValueError, match="Failed to parse DOCX"):
        extract_text(corrupt_docx)
