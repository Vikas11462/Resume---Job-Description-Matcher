"""Document extraction module for Resume-JD Matcher.

Supports text extraction from PDF (.pdf) and Microsoft Word (.docx) documents
with robust validation and error handling.
"""

from pathlib import Path
from typing import Union
import docx
import pdfplumber


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extracts raw textual content from a PDF file using pdfplumber.

    Args:
        pdf_path (Path): Path to the target PDF document.

    Returns:
        str: Extracted text normalized across all document pages.

    Raises:
        ValueError: If the PDF is corrupted or unreadable.
    """
    try:
        pages_text: list[str] = []
        with pdfplumber.open(pdf_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text.strip())
        return "\n\n".join(pages_text).strip()
    except Exception as exc:
        raise ValueError(f"Failed to parse PDF document '{pdf_path.name}': {exc}") from exc


def extract_text_from_docx(docx_path: Path) -> str:
    """Extracts textual content from a DOCX file using python-docx.

    Extracts text from both body paragraphs and embedded table cells.

    Args:
        docx_path (Path): Path to the target DOCX document.

    Returns:
        str: Extracted text normalized across paragraphs and tables.

    Raises:
        ValueError: If the DOCX is corrupted or unreadable.
    """
    try:
        doc = docx.Document(str(docx_path))
        text_chunks: list[str] = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_chunks.append(para.text.strip())

        # Extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_chunks.append(" | ".join(row_text))

        return "\n".join(text_chunks).strip()
    except Exception as exc:
        raise ValueError(f"Failed to parse DOCX document '{docx_path.name}': {exc}") from exc


def extract_text(file_path: Union[str, Path]) -> str:
    """Main extraction dispatcher for resume files.

    Validates file existence and delegates to the appropriate format parser.

    Args:
        file_path (Union[str, Path]): Path to the resume file (.pdf or .docx).

    Returns:
        str: Extracted plain text content.

    Raises:
        FileNotFoundError: If the specified file does not exist.
        ValueError: If the file extension is unsupported or content is corrupt.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if not path.is_file():
        raise ValueError(f"Path is not a valid file: {path}")

    extension = path.suffix.lower()

    if extension == ".pdf":
        return extract_text_from_pdf(path)
    elif extension == ".docx":
        return extract_text_from_docx(path)
    else:
        raise ValueError(
            f"Unsupported file format '{extension}'. Only .pdf and .docx files are supported."
        )
